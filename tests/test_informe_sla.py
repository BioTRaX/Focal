"""Handler para generar informes de SLA."""

from __future__ import annotations

import logging
import os
import tempfile
import locale
from types import SimpleNamespace
from typing import Optional

import pandas as pd
from docx import Document
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import ContextTypes

# ► Exportar a PDF (solo funciona en entornos donde esté disponible)
try:  # pragma: no cover
    import win32com.client as win32  # type: ignore
except Exception:  # pragma: no cover
    win32 = None

import sys
from pathlib import Path

ROOT_DIR = Path(__file__).resolve().parents[1]
sys.path.append(str(ROOT_DIR / "Sandy bot"))

from sandybot.config import config
from sandybot.utils import obtener_mensaje
# Importaciones necesarias solo para las funciones de este test
# Evitamos cargar todo ``sandybot.handlers`` para prevenir errores

# Plantilla predeterminada
RUTA_PLANTILLA = config.SLA_PLANTILLA_PATH

logger = logging.getLogger(__name__)


# ─────────────────────────────── INICIO ──────────────────────────────
async def iniciar_informe_sla(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Coloca al usuario en modo *informe_sla* y solicita los dos Excel."""
    mensaje = obtener_mensaje(update)
    if not mensaje:
        logger.warning("No se recibió mensaje en iniciar_informe_sla")
        return

    user_id = update.effective_user.id
    UserState.set_mode(user_id, "informe_sla")
    context.user_data.clear()
    context.user_data["archivos"] = [None, None]  # [reclamos, servicios]

    # Botón para actualizar plantilla
    try:
        kb = InlineKeyboardMarkup(
            [[InlineKeyboardButton("Actualizar plantilla", callback_data="sla_cambiar_plantilla")]]
        )
    except Exception:  # fallback en tests
        btn = SimpleNamespace(text="Actualizar plantilla", callback_data="sla_cambiar_plantilla")
        kb = SimpleNamespace(inline_keyboard=[[btn]])

    await responder_registrando(
        mensaje,
        user_id,
        "informe_sla",
        "Enviá el Excel de **reclamos** y luego el de **servicios** para generar el informe.",
        "informe_sla",
        reply_markup=kb,
    )


# ────────────────────────── PROCESO COMPLETO ─────────────────────────
async def procesar_informe_sla(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Gestiona la carga de Excel, generación y envío del informe SLA."""
    mensaje = obtener_mensaje(update)
    if not mensaje:
        logger.warning("No se recibió mensaje en procesar_informe_sla")
        return

    user_id = update.effective_user.id
    archivos = context.user_data.setdefault("archivos", [None, None])

    # 1) ── Callback para cambiar plantilla ───────────────────────────
    if update.callback_query and update.callback_query.data == "sla_cambiar_plantilla":
        context.user_data["cambiar_plantilla"] = True
        await update.callback_query.message.reply_text("Adjuntá la nueva plantilla .docx.")
        return

    # 2) ── Guardar la nueva plantilla, si se solicitó ────────────────
    if context.user_data.get("cambiar_plantilla"):
        if getattr(mensaje, "document", None):
            await _actualizar_plantilla_sla(mensaje, context)
        else:
            await responder_registrando(
                mensaje,
                user_id,
                getattr(mensaje, "text", ""),
                "Adjuntá el archivo .docx para actualizar la plantilla.",
                "informe_sla",
            )
        return

    # 3) ── Callback «Procesar informe» ───────────────────────────────
    if update.callback_query and update.callback_query.data == "sla_procesar":
        try:
            ruta_final = _generar_documento_sla(*archivos)
            with open(ruta_final, "rb") as f:
                await update.callback_query.message.reply_document(f, filename=os.path.basename(ruta_final))
            registrar_conversacion(
                user_id, "informe_sla", f"Documento {os.path.basename(ruta_final)} enviado", "informe_sla"
            )
        except Exception as e:  # pragma: no cover
            logger.error("Error generando informe SLA: %s", e)
            await update.callback_query.message.reply_text("💥 Algo falló generando el informe de SLA.")
        finally:
            for p in archivos:
                try:
                    os.remove(p)
                except OSError:
                    pass
            if "ruta_final" in locals() and os.path.exists(ruta_final):
                os.remove(ruta_final)
            context.user_data.clear()
            UserState.set_mode(user_id, "")
        return

    # 4) ── Recepción de archivos Excel ───────────────────────────────
    docs = [d for d in (getattr(mensaje, "document", None), *getattr(mensaje, "documents", [])) if d]
    if docs:
        for doc in docs:
            f = await doc.get_file()
            with tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx") as tmp:
                await f.download_to_drive(tmp.name)
                nombre = doc.file_name.lower()
                if "recl" in nombre and archivos[0] is None:
                    archivos[0] = tmp.name
                elif "serv" in nombre and archivos[1] is None:
                    archivos[1] = tmp.name
                elif archivos[0] is None:
                    archivos[0] = tmp.name
                else:
                    archivos[1] = tmp.name

        if None in archivos:
            falta = "reclamos" if archivos[0] is None else "servicios"
            await responder_registrando(
                mensaje, user_id, docs[-1].file_name,
                f"Archivo guardado. Falta el Excel de {falta}.", "informe_sla",
            )
            return

        # Mostrar botón Procesar
        try:
            kb = InlineKeyboardMarkup(
                [[InlineKeyboardButton("Procesar informe 🚀", callback_data="sla_procesar")]]
            )
        except Exception:  # fallback stubs
            btn = SimpleNamespace(text="Procesar informe 🚀", callback_data="sla_procesar")
            kb = SimpleNamespace(inline_keyboard=[[btn]])

        await responder_registrando(
            mensaje, user_id, docs[-1].file_name,
            "Archivos cargados. Presioná *Procesar informe*.", "informe_sla", reply_markup=kb,
        )
        return

    # 5) ── Ningún adjunto ni callback reconocido ─────────────────────
    await responder_registrando(
        mensaje, user_id, getattr(mensaje, "text", ""),
        "Adjuntá los archivos de reclamos y servicios para comenzar.", "informe_sla",
    )


# ──────────────────── ACTUALIZAR PLANTILLA SLA ───────────────────────
async def _actualizar_plantilla_sla(mensaje, context):
    user_id = mensaje.from_user.id
    archivo = mensaje.document
    if not archivo.file_name.lower().endswith(".docx"):
        await responder_registrando(mensaje, user_id, archivo.file_name, "El archivo debe ser .docx.", "informe_sla")
        return
    try:
        f = await archivo.get_file()
        os.makedirs(os.path.dirname(RUTA_PLANTILLA), exist_ok=True)
        await f.download_to_drive(RUTA_PLANTILLA)
        texto = "Plantilla de SLA actualizada."
        context.user_data.pop("cambiar_plantilla", None)
    except Exception as exc:  # pragma: no cover
        logger.error("Error guardando plantilla SLA: %s", exc)
        texto = "No se pudo guardar la plantilla."

    await responder_registrando(mensaje, user_id, archivo.file_name, texto, "informe_sla")


# ─────────────────── FUNCIÓN GENERADORA DE WORD ──────────────────────
def _generar_documento_sla(
    reclamos_xlsx: str,
    servicios_xlsx: str,
    eventos: Optional[str] = "",
    conclusion: Optional[str] = "",
    propuesta: Optional[str] = "",
    *,
    exportar_pdf: bool = False,
) -> str:
    """Genera el documento SLA; si `exportar_pdf` es True intenta generar PDF."""
    reclamos_df = pd.read_excel(reclamos_xlsx)
    servicios_df = pd.read_excel(servicios_xlsx)

    columnas_extra = [c for c in ("SLA Entregado", "Dirección", "Horas Netas Reclamo") if c in servicios_df]

    # Normaliza nombres
    if "Servicio" not in reclamos_df.columns:
        reclamos_df.rename(columns={reclamos_df.columns[0]: "Servicio"}, inplace=True)
    if "Servicio" not in servicios_df.columns:
        servicios_df.rename(columns={servicios_df.columns[0]: "Servicio"}, inplace=True)

    # Fecha para título
    try:
        fecha = pd.to_datetime(reclamos_df.iloc[0].get("Fecha"))
        if pd.isna(fecha):
            raise ValueError
    except Exception:
        fecha = pd.Timestamp.today()

    # Locale español (ignora errores si no está instalado)
    for loc in ("es_ES.UTF-8", "es_ES", "es_AR.UTF-8", "es_AR"):
        try:
            locale.setlocale(locale.LC_TIME, loc)
            break
        except locale.Error:
            continue

    mes, anio = fecha.strftime("%B").upper(), fecha.strftime("%Y")

    # Conteo de reclamos
    resumen = reclamos_df.groupby("Servicio").size().reset_index(name="Reclamos")
    df = servicios_df.merge(resumen, on="Servicio", how="left")
    df["Reclamos"] = df["Reclamos"].fillna(0).astype(int)

    # Documento base
    if not (RUTA_PLANTILLA and os.path.exists(RUTA_PLANTILLA)):
        raise ValueError(f"Plantilla de SLA no encontrada: {RUTA_PLANTILLA}")
    doc = Document(RUTA_PLANTILLA)

    try:
        doc.add_heading(f"Informe SLA {mes} {anio}", level=0)
    except KeyError:
        doc.add_heading(f"Informe SLA {mes} {anio}", level=1)

    # Tabla resumen
    headers = ["Servicio", *columnas_extra, "Reclamos"]
    tbl = doc.add_table(rows=1, cols=len(headers), style="Table Grid")
    for i, h in enumerate(headers):
        tbl.rows[0].cells[i].text = h

    for _, fila in df.iteritems():
        cells = tbl.add_row().cells
        for i, h in enumerate(headers):
            cells[i].text = str(fila.get(h, ""))

    # Etiquetas dinámicas
    etiquetas = {
        "Eventos sucedidos de mayor impacto en SLA:": eventos,
        "Conclusión:": conclusion,
        "Propuesta de mejora:": propuesta,
    }
    encontrados = set()
    for p in doc.paragraphs:
        for etq, cont in etiquetas.items():
            if p.text.strip().startswith(etq):
                p.text = f"{etq} {cont}"
                encontrados.add(etq)
                break
    for etq, cont in etiquetas.items():
        if etq not in encontrados and cont:
            doc.add_paragraph(f"{etq} {cont}")

    # Guardar DOCX
    fd, ruta_docx = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    doc.save(ruta_docx)

    # Exportar PDF (opcional)
    if exportar_pdf:
        pdf_path = os.path.splitext(ruta_docx)[0] + ".pdf"
        convertido = False

        if win32 and os.name == "nt":
            try:
                word = win32.Dispatch("Word.Application")
                word_doc = word.Documents.Open(ruta_docx)
                word_doc.SaveAs(pdf_path, FileFormat=17)
                word_doc.Close()
                word.Quit()
                convertido = True
            except Exception:
                logger.warning("Fallo exportando PDF con win32")

        if not convertido:
            try:
                from docx2pdf import convert  # type: ignore
                convert(ruta_docx, pdf_path)
                convertido = True
            except Exception:
                logger.warning("Fallo exportando PDF con docx2pdf")

        if convertido:
            os.remove(ruta_docx)
            return pdf_path

    return ruta_docx


def test_documento_sla_columna_faltante(tmp_path, caplog):
    """Genera el informe aun cuando falta una columna y muestra advertencia."""
    import sys
    import importlib
    from pathlib import Path

    ROOT_DIR = Path(__file__).resolve().parents[1]
    sys.path.append(str(ROOT_DIR / "Sandy bot"))

    # Plantilla de prueba
    plantilla = tmp_path / "tpl.docx"
    doc = Document()
    doc.add_paragraph("Eventos sucedidos de mayor impacto en SLA:")
    doc.add_paragraph("Conclusión:")
    doc.add_paragraph("Propuesta de mejora:")
    doc.save(plantilla)
    os.environ["SLA_TEMPLATE_PATH"] = str(plantilla)

    # Variables minimas para Config
    for var in [
        "TELEGRAM_TOKEN",
        "OPENAI_API_KEY",
        "NOTION_TOKEN",
        "NOTION_DATABASE_ID",
        "DB_USER",
        "DB_PASSWORD",
        "SLACK_WEBHOOK_URL",
        "SUPERVISOR_DB_ID",
    ]:
        os.environ.setdefault(var, "x")

    # Recargar módulos para que tomen la plantilla recién definida
    config_mod = importlib.reload(importlib.import_module("sandybot.config"))
    informe = importlib.reload(importlib.import_module("sandybot.handlers.informe_sla"))

    reclamos = tmp_path / "reclamos.xlsx"
    servicios = tmp_path / "servicios.xlsx"
    pd.DataFrame({"Servicio": ["S1"], "Fecha": ["2024-01-01"]}).to_excel(reclamos, index=False)
    pd.DataFrame({
        "Servicio": ["S1"],
        "SLA Entregado": [95],
        # Falta la columna "Dirección"
        "Horas Netas Reclamo": [2],
    }).to_excel(servicios, index=False)

    with caplog.at_level(logging.WARNING):
        ruta = informe._generar_documento_sla(str(reclamos), str(servicios))

    assert os.path.exists(ruta)
    assert "Faltan columnas" in caplog.text
    os.remove(ruta)
