# Nombre de archivo: repetitividad.py
# Ubicación de archivo: Sandy bot/sandybot/handlers/repetitividad.py
# User-provided custom instructions
"""
Handler para la generación de informes de repetitividad.
"""
from telegram import Update
from telegram.ext import ContextTypes
import os
import tempfile
import pandas as pd
from docx import Document
# Los siguientes módulos solo están disponibles en Windows. Se usan para
# modificar el documento Word mediante COM.
try:
    import win32com.client as win32
    import pythoncom
except ImportError:
    # Si fallan las importaciones, se asigna None para deshabilitar
    # la funcionalidad específica de Windows.
    win32 = None
    pythoncom = None
import locale
from datetime import datetime
import logging

from sandybot.config import config
from ..utils import obtener_mensaje
from .estado import UserState
from ..registrador import responder_registrando, registrar_conversacion

# Ruta a la plantilla Word definida en la configuración global
# Permite modificar la ubicación mediante la variable de entorno "PLANTILLA_PATH"
RUTA_PLANTILLA = config.PLANTILLA_PATH

logger = logging.getLogger(__name__)

async def manejar_repetitividad(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """
    Maneja la generación de informes de repetitividad.

    :param update: Objeto de actualización de Telegram.
    :param context: Contexto del manejador.
    """
    message = obtener_mensaje(update)
    if not message:
        logger.warning("No se obtuvo mensaje en manejar_repetitividad.")
        return

    try:
        logger.info(
            "Iniciando manejo de repetitividad para el usuario %s",
            update.effective_user.id,
        )
        await responder_registrando(
            message,
            update.effective_user.id,
            "informe_repetitividad",
            "Generación de informes de repetitividad en desarrollo.",
            "repetitividad",
        )
    except Exception as e:
        logger.error("Error en manejar_repetitividad: %s", e)
        if message:
            await responder_registrando(
                message,
                update.effective_user.id,
                "informe_repetitividad",
                f"Error al generar el informe: {e}",
                "repetitividad",
            )

async def iniciar_repetitividad(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """
    Inicia el proceso de generación de informes de repetitividad.

    :param update: Objeto de actualización de Telegram.
    :param context: Contexto del manejador.
    """
    message = obtener_mensaje(update)
    if not message:
        logger.warning("No se obtuvo mensaje en iniciar_repetitividad.")
        return

    try:
        logger.info(
            "Iniciando repetitividad para el usuario %s",
            update.effective_user.id,
        )
        await responder_registrando(
            message,
            update.effective_user.id,
            "informe_repetitividad",
            "Iniciando generación de informes de repetitividad. Enviá el archivo Excel para continuar.",
            "repetitividad",
        )
    except Exception as e:
        logger.error("Error en iniciar_repetitividad: %s", e)
        if message:
            await responder_registrando(
                message,
                update.effective_user.id,
                "informe_repetitividad",
                f"Error al iniciar la generación de informes: {e}",
                "repetitividad",
            )

async def procesar_repetitividad(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """
    Procesa los datos para generar un informe de repetitividad.

    :param update: Objeto de actualización de Telegram.
    :param context: Contexto del manejador.
    """
    message = obtener_mensaje(update)
    if not message:
        logger.warning("No se obtuvo mensaje en procesar_repetitividad.")
        return
    user_id = message.from_user.id

    try:
        if not message.document:
            await responder_registrando(
                message,
                user_id,
                "procesar_repetitividad",
                "😠 ¿Y el archivo? Adjuntá el Excel, por favor. No soy adivino. #DaleCerebro",
                "repetitividad",
            )
            return

        archivo = message.document
        if not archivo.file_name.endswith(".xlsx"):
            await responder_registrando(
                message,
                user_id,
                archivo.file_name,
                "🙄 Solo acepto archivos Excel (.xlsx). Mandá bien las cosas. #MeEstásCargando",
                "repetitividad",
            )
            return

        file = await archivo.get_file()
        with tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx") as tmp_excel:
            await file.download_to_drive(tmp_excel.name)

        try:
            ruta_salida = generar_informe_y_modificar(tmp_excel.name)
        except ValueError as err:
            await responder_registrando(
                message,
                user_id,
                archivo.file_name,
                str(err),
                "repetitividad",
            )
            return
        nombre_final = os.path.basename(ruta_salida)
        with open(ruta_salida, "rb") as docx_file:
            await message.reply_document(document=docx_file, filename=nombre_final)
        registrar_conversacion(
            user_id,
            archivo.file_name,
            f"Documento {nombre_final} enviado",
            "repetitividad",
        )

        os.remove(tmp_excel.name)
        os.remove(ruta_salida)
        UserState.set_mode(user_id, "")

    except Exception as e:
        if message:
            await responder_registrando(
                message,
                user_id,
                archivo.file_name if 'archivo' in locals() else 'procesar_repetitividad',
                "💥 Algo falló generando el informe. No sé cómo hacés para romper hasta esto... #LoQueHayQueAguantar",
                "repetitividad",
            )


def generar_informe_y_modificar(ruta_excel):
    for loc in ("es_ES.UTF-8", "es_ES", "es_AR.UTF-8", "es_AR"):
        try:
            locale.setlocale(locale.LC_TIME, loc)
            break
        except locale.Error:
            continue

    try:
        casos_df = pd.read_excel(ruta_excel)
    except Exception as exc:
        logger.error("Error leyendo el Excel %s: %s", ruta_excel, exc)
        raise ValueError("⚠️ No se pudo leer el Excel. Verificá el archivo.") from exc

    columnas_a_conservar_casos = [
        'Número Reclamo',
        'Número Línea',
        'Tipo Servicio',
        'Nombre Cliente',
        'Fecha Inicio Reclamo',
        'Fecha Cierre Reclamo',
        'Fecha Cierre Problema Reclamo',
        'Horas Netas Problema Reclamo',
        'Tipo Solución Reclamo',
        'Descripción Solución Reclamo',
    ]

    faltantes = set(columnas_a_conservar_casos) - set(casos_df.columns)
    if faltantes:
        logger.error("Faltan columnas requeridas: %s", ", ".join(faltantes))
        raise ValueError(
            "⚠️ El Excel no tiene todas las columnas necesarias. Revisá el formato."
        )

    casos_limpio = casos_df[columnas_a_conservar_casos].copy()

    fecha_cierre = None
    for col in ['Fecha Cierre Reclamo', 'Fecha Cierre Problema Reclamo']:
        try:
            valor_fecha = casos_limpio[col].dropna().iloc[0]
            fecha_cierre = pd.to_datetime(valor_fecha)
            break
        except Exception:
            continue
    if fecha_cierre is None:
        logger.warning(
            "No se pudieron leer las fechas de cierre, se usa la actual",
        )
        fecha_cierre = datetime.now()

    mes_actual = fecha_cierre.strftime("%B")
    año_actual = fecha_cierre.strftime("%Y")

    casos_limpio['Número Línea'] = casos_limpio['Número Línea'].astype(str).str.replace('.0', '', regex=False)
    casos_limpio = casos_limpio.sort_values(by='Número Línea')
    lineas_con_multiples_reclamos = casos_limpio['Número Línea'].value_counts()
    lineas_a_conservar = lineas_con_multiples_reclamos[lineas_con_multiples_reclamos >= 2].index
    casos_filtrados = casos_limpio[casos_limpio['Número Línea'].isin(lineas_a_conservar)]

    if not os.path.exists(RUTA_PLANTILLA):
        raise ValueError(
            f"⚠️ No se encontró la plantilla en {RUTA_PLANTILLA}. \
Configurá la variable PLANTILLA_PATH."
        )
    doc = Document(RUTA_PLANTILLA)

    for numero_linea, grupo in casos_filtrados.groupby('Número Línea'):
        nombre_cliente = grupo['Nombre Cliente'].iloc[0]
        tipo_servicio = grupo['Tipo Servicio'].iloc[0]
        doc.add_paragraph(f'{tipo_servicio}: {numero_linea} - {nombre_cliente}', style='Heading 1')

        tabla = doc.add_table(rows=1, cols=7, style='Table Grid')
        hdr_cells = tabla.rows[0].cells
        hdr_cells[0].text = 'Reclamo'
        hdr_cells[1].text = 'Tipo Solución Reclamo'
        hdr_cells[2].text = 'Fecha Inicio Reclamo'
        hdr_cells[3].text = 'Fecha Cierre Reclamo'
        hdr_cells[4].text = 'Fecha Cierre Problema Reclamo'
        hdr_cells[5].text = 'Horas Netas Problema Reclamo'
        hdr_cells[6].text = 'Descripción Solución Reclamo'

        for _, fila in grupo.iterrows():
            fila_cells = tabla.add_row().cells
            fila_cells[0].text = str(fila['Número Reclamo'])
            fila_cells[1].text = fila['Tipo Solución Reclamo']
            fila_cells[2].text = (
                fila['Fecha Inicio Reclamo'].strftime('%d/%m/%Y %H:%M')
                if pd.notnull(fila['Fecha Inicio Reclamo'])
                else ''
            )
            fila_cells[3].text = (
                fila['Fecha Cierre Reclamo'].strftime('%d/%m/%Y %H:%M')
                if pd.notnull(fila['Fecha Cierre Reclamo'])
                else ''
            )
            fila_cells[4].text = (
                fila['Fecha Cierre Problema Reclamo'].strftime('%d/%m/%Y %H:%M')
                if pd.notnull(fila['Fecha Cierre Problema Reclamo'])
                else ''
            )
            if pd.notnull(fila['Horas Netas Problema Reclamo']):
                horas_valor = fila['Horas Netas Problema Reclamo']
                if isinstance(horas_valor, pd.Timedelta):
                    total_min = int(horas_valor.total_seconds() // 60)
                    horas = total_min // 60
                    minutos = total_min % 60
                    fila_cells[5].text = f"{horas:02d}:{minutos:02d} Hrs"
                else:
                    fila_cells[5].text = str(horas_valor)
            else:
                fila_cells[5].text = ''
            fila_cells[6].text = fila['Descripción Solución Reclamo']

    nombre_archivo = f"InformeRepetitividad{fecha_cierre.strftime('%m%y')}.docx"
    ruta_docx_generado = os.path.join(tempfile.gettempdir(), nombre_archivo)
    doc.save(ruta_docx_generado)

    # Si win32 es None significa que la importación falló y estamos en un
    # sistema que no soporta esta funcionalidad (generalmente no Windows).
    if win32 is not None:
        modificar_informe_con_pythoncom(ruta_docx_generado, mes_actual, año_actual)
    else:
        logger.info(
            "Omitiendo modificación por COM; esta funcionalidad solo está disponible en Windows."
        )
    return ruta_docx_generado


def modificar_informe_con_pythoncom(docx_path, mes_actual, año_actual):
    """Aplica cambios en el Word mediante COM. Solo disponible en Windows."""
    pythoncom.CoInitialize()
    try:
        word_app = win32.Dispatch("Word.Application")
        word_app.Visible = False
        doc = word_app.Documents.Open(docx_path)

        titulo_dinamico = f"Informe Repetitividad {mes_actual} {año_actual}"

        for shape in doc.Shapes:
            if shape.TextFrame.HasText and "Informe" in shape.TextFrame.TextRange.Text:
                shape.TextFrame.TextRange.Text = titulo_dinamico

        for tabla in doc.Tables:
            for cell in tabla.Rows(1).Cells:
                cell.Shading.BackgroundPatternColor = 12611584

        doc.SaveAs(docx_path)
        doc.Close()
        word_app.Quit()
    except Exception as e:
        logger.error("Error al modificar el documento con COM: %s", e)
    finally:
        pythoncom.CoUninitialize()
